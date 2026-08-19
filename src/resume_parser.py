import os
import json
import re
from pypdf import PdfReader

try:
    import docx
except ImportError:
    docx = None

from google import genai
from src.config import Config

def extract_resume_text(resume_path: str) -> str:
    """
    Extracts text content from either a .pdf or .docx resume file.
    """
    if not os.path.exists(resume_path):
        raise FileNotFoundError(f"Resume file not found at: {resume_path}")

    ext = os.path.splitext(resume_path)[1].lower()

    if ext == ".pdf":
        reader = PdfReader(resume_path)
        text = ""
        for page in reader.pages:
            content = page.extract_text()
            if content:
                text += content + "\n"
        return text.strip()

    elif ext in [".docx", ".doc"]:
        if docx is None:
            raise ImportError("python-docx package is required to read .docx files. Install with 'pip install python-docx'.")
        
        doc = docx.Document(resume_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return "\n".join(full_text).strip()

    else:
        raise ValueError(f"Unsupported file format '{ext}'. Please provide a .pdf or .docx file.")

def extract_relevant_roles_from_resume(resume_text: str, api_key: str = None) -> dict:
    """
    Uses Gemini AI to intelligently analyze any candidate resume (Software, Data, Design, Marketing, Finance, etc.)
    and automatically extract:
    1. Primary Job Role
    2. 3-5 Tailored Job Search Queries / Job Titles
    3. Core Skills list
    """
    api_key = api_key or os.getenv("GEMINI_API_KEY", Config.GEMINI_API_KEY)
    
    if api_key:
        try:
            client = genai.Client(api_key=api_key)
            prompt = f"""
            You are an expert AI Career Recruiter and Job Matching Engine.
            Analyze the following resume and determine what jobs this candidate is qualified to apply for.
            
            Resume Text:
            {resume_text[:3000]}
            
            Return ONLY a valid JSON object with the following structure:
            {{
                "primary_role": "<The most suitable primary job title for this resume>",
                "target_roles": ["<Role 1>", "<Role 2>", "<Role 3>", "<Role 4>"],
                "core_skills": ["<Skill 1>", "<Skill 2>", "<Skill 3>"],
                "summary": "<1-2 sentence professional summary>"
            }}
            Do NOT include markdown wrapping outside the JSON.
            """
            
            response = client.models.generate_content(
                model="gemini-3.6-flash",
                contents=prompt
            )
            raw = response.text.replace("```json", "").replace("```", "").strip()
            data = json.loads(raw)
            if "primary_role" in data and "target_roles" in data:
                return data
        except Exception as e:
            print(f"[WARNING] Gemini resume analysis error: {e}")

    # Heuristic fallback if AI unavailable
    first_lines = resume_text[:500].lower()
    if "developer" in first_lines or "software" in first_lines or "react" in first_lines:
        return {
            "primary_role": "Software Developer",
            "target_roles": ["Full Stack Developer", "Frontend Developer", "Node.js Developer", "React Developer"],
            "core_skills": ["JavaScript", "React", "Node.js"],
            "summary": "Experienced software developer."
        }
    elif "data" in first_lines or "analytics" in first_lines:
        return {
            "primary_role": "Data Analyst",
            "target_roles": ["Data Analyst", "Data Engineer", "Business Intelligence Analyst", "Python Data Specialist"],
            "core_skills": ["SQL", "Python", "Data Analysis"],
            "summary": "Data specialist with analytical expertise."
        }
    elif "design" in first_lines or "ui" in first_lines or "ux" in first_lines:
        return {
            "primary_role": "UI/UX Designer",
            "target_roles": ["UI/UX Designer", "Product Designer", "Web Designer", "Graphic Designer"],
            "core_skills": ["Figma", "UI/UX", "Design Systems"],
            "summary": "Creative designer specializing in user experience."
        }
    
    return {
        "primary_role": "Professional Candidate",
        "target_roles": ["Operations Specialist", "Project Coordinator", "Associate Specialist"],
        "core_skills": ["Management", "Communication"],
        "summary": "Versatile candidate matching multiple business roles."
    }
